import pandas as pd
from docx import Document
import sys
import os

def gerar_memorial_descritivo(caminho_planilha, caminho_modelo, caminho_saida):
    """
    Gera um memorial descritivo em formato de texto a partir de uma planilha de dados.

    Args:
        caminho_planilha (str): Caminho para o arquivo de dados.
        caminho_modelo (str): Caminho para o modelo do memorial.
        caminho_saida (str): Caminho para salvar o documento final.
    """
    try:
        df = pd.read_excel(caminho_planilha)
    except FileNotFoundError:
        print(f"Erro: O arquivo de planilha não foi encontrado em {caminho_planilha}")
        return

    # Extração das informações gerais (buscando cada uma na sua devida linha)
    nome_imovel = df.loc[0, 'Nome Imovel']
    proprietario = df.loc[4, 'Proprietario']
    area = df.loc[0, 'Area']
    matricula = df.loc[2, 'Matricula']
    perimetro = df.loc[2, 'Perimetro']
    municipio = df.loc[0, 'Municipio']
    uf = df.loc[0, 'UF']
    comarca = df.loc[2, 'Comarca']
    cpf = df.loc[4, 'CPF']
    trt = df.loc[2, 'TRT']

    # Geração do Documento com python-docx
    doc = Document(caminho_modelo)

    # Substitui os placeholders no documento
    for p in doc.paragraphs:
        if '[NOME_IMOVEL]' in p.text:
            p.text = p.text.replace('[NOME_IMOVEL]', str(nome_imovel))
        if '[PROPRIETARIO]' in p.text:
            p.text = p.text.replace('[PROPRIETARIO]', str(proprietario))
        if '[AREA]' in p.text:
            p.text = p.text.replace('[AREA]', str(area))
        if '[MATRICULA]' in p.text:
            p.text = p.text.replace('[MATRICULA]', str(matricula))
        if '[PERIMETRO]' in p.text:
            p.text = p.text.replace('[PERIMETRO]', str(perimetro))
        if '[MUNICIPIO]' in p.text:
            p.text = p.text.replace('[MUNICIPIO]', str(municipio))
        if '[UF]' in p.text:
            p.text = p.text.replace('[UF]', str(uf))
        if '[COMARCA]' in p.text:
            p.text = p.text.replace('[COMARCA]', str(comarca))
        if '[CPF]' in p.text:
            p.text = p.text.replace('[CPF]', str(cpf))
        if '[TRT]' in p.text:
            p.text = p.text.replace('[TRT]', str(trt))

    # Adiciona a descrição do perímetro em parágrafos
    doc.add_paragraph()
    doc.add_paragraph("DESCRIÇÃO")

    # Inserindo o primeiro parágrafo de descrição (início da descrição)
    primeiro_vertice = df.loc[0, 'Vertice']
    coord_N_primeiro = df.loc[0, 'coord_N']
    coord_E_primeiro = df.loc[0, 'coord_E']
    confrontante_primeiro = df.loc[0, 'Confrontante']
    
    doc.add_paragraph(f"Inicia-se a descrição deste perímetro no vértice {primeiro_vertice}, de coordenadas N {coord_N_primeiro}m e E {coord_E_primeiro}m; situado no limite com {confrontante_primeiro};")

    # Itera sobre as linhas da planilha para criar os parágrafos de cada trecho
    for index, row in df.iterrows():
        azimute_atual = row['Azimute']
        distancia_atual = row['Distancia']
        confrontante_atual = row['Confrontante']
        
        # O próximo vértice, se houver
        if index + 1 < len(df):
            proximo_vertice = df.loc[index + 1, 'Vertice']
            coord_N_proximo = df.loc[index + 1, 'coord_N']
            coord_E_proximo = df.loc[index + 1, 'coord_E']
            desc_paragrafo = f"deste, segue com azimute e distância de: {azimute_atual} e {distancia_atual}m, confrontando neste trecho com {confrontante_atual} até o vértice {proximo_vertice}, de coordenadas N {coord_N_proximo}m e E {coord_E_proximo}m;"
        else:
            desc_paragrafo = f"deste, segue com azimute e distância de: {azimute_atual} e {distancia_atual}m, confrontando neste trecho com {confrontante_atual} até o vértice {primeiro_vertice}, ponto inicial da descrição deste perímetro."

        doc.add_paragraph(desc_paragrafo)

    doc.add_paragraph("Todas as coordenadas aqui descritas estão demarcadas e encontram-se representadas no Sistema UTM (Universal Transversa de Mercator), referenciadas ao Meridiano Central 45° W (Fuso 23), tendo como o Datum o SIRGAS-2000.")

    # Salva o documento final
    doc.save(caminho_saida)
    print(f"Memorial descritivo gerado com sucesso em '{caminho_saida}'.")

# --- Lógica para o modo de arrastar e soltar ---
# Verifica se o script recebeu um argumento (caminho do arquivo)
if len(sys.argv) > 1:
    caminho_da_planilha = sys.argv[1]
    # Pega o caminho do diretório onde o script e o modelo estão
    diretorio_script = os.path.dirname(os.path.abspath(__file__))
    caminho_do_modelo = os.path.join(diretorio_script, "modelo_memorial.docx")
    caminho_de_saida = os.path.join(diretorio_script, "memorial_final.docx")
    
    gerar_memorial_descritivo(caminho_da_planilha, caminho_do_modelo, caminho_de_saida)
else:
    print("Por favor, arraste e solte o arquivo da planilha (dados_memorial.xlsx) no script.")